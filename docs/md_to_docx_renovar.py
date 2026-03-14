#!/usr/bin/env python3
"""
Convert Renovar_Startup_Business_Plan.md to a Word document.
Requires: pip install python-docx
"""
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

SCRIPT_DIR = Path(__file__).resolve().parent
MD_PATH = SCRIPT_DIR / "Renovar_Startup_Business_Plan.md"
OUTPUT_DOCX = SCRIPT_DIR / "Renovar_Startup_Business_Plan.docx"


def parse_table(lines, start_idx):
    """Parse markdown table; return (list of rows, next_index)."""
    rows = []
    i = start_idx
    while i < len(lines):
        line = lines[i]
        if "|" not in line or not line.strip():
            break
        # Skip separator row like |---|---|
        if re.match(r"^\s*\|[\s\-:]+\|", line):
            i += 1
            continue
        parts = line.split("|")
        cells = [c.strip() for c in parts[1:-1]] if len(parts) > 2 else [c.strip() for c in parts]
        if cells:
            rows.append(cells)
        i += 1
    return rows, i


def add_bold_runs(paragraph, text):
    """Add text to paragraph with **bold** rendered as bold."""
    rest = text
    while rest:
        m = re.search(r"\*\*(.+?)\*\*", rest)
        if not m:
            paragraph.add_run(rest)
            break
        paragraph.add_run(rest[: m.start()])
        run = paragraph.add_run(m.group(1))
        run.bold = True
        rest = rest[m.end() :]
    return paragraph


def md_to_docx():
    doc = Document()
    text = MD_PATH.read_text(encoding="utf-8")
    lines = text.split("\n")

    i = 0
    in_code_block = False
    code_lines = []

    while i < len(lines):
        line = lines[i]
        raw = line

        # Code block
        if line.strip().startswith("```"):
            if in_code_block:
                p = doc.add_paragraph("\n".join(code_lines))
                p.style = "Normal"
                for r in p.runs:
                    r.font.name = "Consolas"
                    r.font.size = Pt(9)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Horizontal rule / empty
        if line.strip() == "---" or line.strip() == "":
            if line.strip() == "---":
                doc.add_paragraph()
            i += 1
            continue

        # Headings
        if line.startswith("# "):
            p = doc.add_heading(line[2:].strip(), level=0)
            i += 1
            continue
        if line.startswith("## "):
            p = doc.add_heading(line[3:].strip(), level=1)
            i += 1
            continue
        if line.startswith("### "):
            p = doc.add_heading(line[4:].strip(), level=2)
            i += 1
            continue

        # Table
        if "|" in line and line.strip().startswith("|"):
            table_rows, next_i = parse_table(lines, i)
            i = next_i
            if table_rows:
                ncols = max(len(r) for r in table_rows)
                table = doc.add_table(rows=len(table_rows), cols=ncols)
                table.style = "Table Grid"
                for ri, row in enumerate(table_rows):
                    for ci, cell in enumerate(row):
                        if ci < ncols:
                            table.rows[ri].cells[ci].text = cell
            continue

        # Bullet
        if line.strip().startswith("- ") or line.strip().startswith("* "):
            content = line.strip()[2:].strip()
            p = doc.add_paragraph(style="List Bullet")
            add_bold_runs(p, content)
            i += 1
            continue

        # Italic-only line
        if line.strip().startswith("*") and line.strip().endswith("*") and " " in line.strip():
            p = doc.add_paragraph()
            run = p.add_run(line.strip()[1:-1])
            run.italic = True
            i += 1
            continue

        # Numbered list (e.g. "1. ")
        if re.match(r"^\s*\d+\.\s", line):
            content = re.sub(r"^\s*\d+\.\s", "", line).strip()
            p = doc.add_paragraph(style="List Number")
            add_bold_runs(p, content)
            i += 1
            continue

        # Normal paragraph (may contain **bold**)
        stripped = line.strip()
        if stripped:
            p = doc.add_paragraph()
            add_bold_runs(p, stripped)
        i += 1

    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_DOCX))
    print(f"Saved: {OUTPUT_DOCX}")


if __name__ == "__main__":
    md_to_docx()
