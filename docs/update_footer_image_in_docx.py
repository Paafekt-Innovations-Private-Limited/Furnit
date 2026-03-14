#!/usr/bin/env python3
"""
Replace the footer image in a letterhead .docx with docs/letterhead/bottom.png.
Usage: python docs/update_footer_image_in_docx.py <path-to-letterhead.docx>
"""
import sys
from pathlib import Path

DOCS_DIR = Path(__file__).resolve().parent
BOTTOM_PNG = DOCS_DIR / "letterhead" / "bottom.png"


def main() -> None:
    if len(sys.argv) < 2:
        raise SystemExit("Usage: python docs/update_footer_image_in_docx.py <path-to-letterhead.docx>")
    docx_path = Path(sys.argv[1]).expanduser().resolve()
    if not docx_path.exists():
        raise SystemExit(f"Not found: {docx_path}")
    if not BOTTOM_PNG.exists():
        raise SystemExit(f"Not found: {BOTTOM_PNG}")

    from docx import Document
    from docx.shared import Inches

    doc = Document(str(docx_path))
    footer_width = Inches(6.5)

    for section in doc.sections:
        footer = section.footer
        if not footer.paragraphs:
            continue
        first = footer.paragraphs[0]
        # Remove all existing runs (old image), then add new image
        for r in list(first.runs):
            first._p.remove(r._element)
        run = first.add_run()
        run.add_picture(str(BOTTOM_PNG), width=footer_width)

    doc.save(str(docx_path))
    print(f"Updated footer image in {docx_path}")


if __name__ == "__main__":
    main()
