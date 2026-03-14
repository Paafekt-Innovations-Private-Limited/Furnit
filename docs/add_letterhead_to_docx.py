#!/usr/bin/env python3
"""
Add Paafekt letterhead (top and bottom bands) to an existing .docx.

Uses the same top/bottom PNGs as the PDF generator:
  docs/letterhead/top.png
  docs/letterhead/bottom.png

Usage (example):
  python docs/add_letterhead_to_docx.py \
      --input "/Users/al/Downloads/NDA - Paafekt v02_0903 (1).docx"

By default writes alongside the input as:
  "<name> - letterhead.docx"
You can override with --output.
"""

from __future__ import annotations

import argparse
from pathlib import Path


DOCS_DIR = Path(__file__).resolve().parent
LETTERHEAD_TOP = DOCS_DIR / "letterhead" / "top.png"
LETTERHEAD_BOTTOM = DOCS_DIR / "letterhead" / "bottom.png"


# Jayamma's Director Identification Number (shown in footer so it is not cut off)
DIN_TEXT = "Director (DIN: 11594139)"


def add_letterhead(input_path: Path, output_path: Path) -> None:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    if not input_path.exists():
        raise FileNotFoundError(f"Input .docx not found: {input_path}")

    if not LETTERHEAD_TOP.exists():
        raise FileNotFoundError(f"Top letterhead image not found: {LETTERHEAD_TOP}")
    if not LETTERHEAD_BOTTOM.exists():
        raise FileNotFoundError(f"Bottom letterhead image not found: {LETTERHEAD_BOTTOM}")

    doc = Document(str(input_path))

    # Heuristic width: leave ~0.75\" side margins on A4/Letter
    header_width = Inches(6.5)
    footer_width = Inches(6.5)

    for section in doc.sections:
        # Header: insert top.png
        header = section.header
        if header.paragraphs:
            h_para = header.paragraphs[0]
        else:
            h_para = header.add_paragraph()
        h_run = h_para.add_run()
        h_run.add_picture(str(LETTERHEAD_TOP), width=header_width)

        # Footer: bottom.png then DIN line (so DIN is visible and not cut)
        footer = section.footer
        if footer.paragraphs:
            f_para = footer.paragraphs[0]
        else:
            f_para = footer.add_paragraph()
        f_run = f_para.add_run()
        f_run.add_picture(str(LETTERHEAD_BOTTOM), width=footer_width)

        din_para = footer.add_paragraph()
        din_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        din_run = din_para.add_run(DIN_TEXT)
        din_run.font.size = Pt(9)
        din_run.font.name = "Helvetica"
        # Tight spacing so image + DIN fit in footer
        f_para.paragraph_format.space_after = Pt(2)
        din_para.paragraph_format.space_before = Pt(0)
        din_para.paragraph_format.space_after = Pt(0)
        # Footer distance = from page bottom to footer bottom; use enough so nothing is cut when printing
        section.footer_distance = Inches(1.0)
        # Keep body (signature block) above the footer so it is not overlapped or cut
        section.bottom_margin = Inches(1.25)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def main() -> None:
    parser = argparse.ArgumentParser(description="Add Paafekt letterhead to a .docx using docs/letterhead images.")
    parser.add_argument("--input", required=True, help="Path to input .docx")
    parser.add_argument("--output", help="Path to output .docx (default: '<name> - letterhead.docx' next to input)")
    args = parser.parse_args()

    in_path = Path(args.input).expanduser().resolve()
    if args.output:
        out_path = Path(args.output).expanduser().resolve()
    else:
        out_path = in_path.with_name(in_path.stem + " - letterhead.docx")

    add_letterhead(in_path, out_path)
    print(f"Wrote letterhead .docx to: {out_path}")


if __name__ == "__main__":
    main()

