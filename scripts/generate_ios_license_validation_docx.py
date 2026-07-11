#!/usr/bin/env python3
"""Generate a non-technical iOS license validation DOCX for external review."""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

REPO_ROOT = Path(__file__).resolve().parents[1]
OUT_PATH = REPO_ROOT / "docs" / "iOS_License_Validation_Checklist.docx"


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value
    doc.add_paragraph()


def main() -> None:
    doc = Document()

    title = doc.add_heading("Paafekt (Furnit) — iOS License Checklist", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Date: {date.today().isoformat()}")
    doc.add_paragraph(
        "Who this is for: a non-technical reviewer. "
        "You do not need the codebase — open the app (or read the text below) and check public license pages on the web."
    )
    doc.add_paragraph(
        "This is not legal advice. When unsure, mark “Ask a lawyer.”"
    )

    doc.add_heading("Where to look in the app", level=1)
    doc.add_paragraph(
        "On iPhone: open Paafekt → Profile (or Settings) → General → scroll to Legal → tap Licenses."
    )
    doc.add_paragraph(
        "You should see a short notice at the top, then a list of software names with descriptions and “View full license” links."
    )

    doc.add_heading("What the app says (copy for your notes)", level=1)

    doc.add_heading("Notice at top of Licenses screen", level=2)
    doc.add_paragraph(
        "“Current release: This app is currently offered for non-commercial use only. "
        "A future commercial release may be subject to different terms.”"
    )
    doc.add_paragraph(
        "This is the app’s own policy text — not a standard open-source license. "
        "Flag if the company plans a normal paid App Store launch."
    )

    doc.add_heading("Open-source items listed in the app", level=2)
    items = [
        (
            "Depth Anything V2",
            "Metric indoor depth estimation for single-photo room reconstruction. "
            "Licensed under the Apache License 2.0.",
            "Apache 2.0",
        ),
        (
            "GeoCalib",
            "Single-image camera calibration for focal length and gravity hints. "
            "Copyright 2024 ETH Zurich. Licensed under the Apache License 2.0.",
            "Apache 2.0",
        ),
        (
            "MetalSplatter",
            "On-device Gaussian splat rendering. Copyright (c) Sean Cier. Licensed under the MIT License.",
            "MIT",
        ),
        (
            "Firebase",
            "Copyright (c) Google LLC. Licensed under the Apache License 2.0.",
            "Apache 2.0",
        ),
        (
            "RTMDet (MMDetection)",
            "Furniture detection and segmentation use the RTMDet-Ins model from OpenMMLab MMDetection. "
            "Copyright (c) OpenMMLab. Licensed under the Apache License 2.0.",
            "Apache 2.0",
        ),
        (
            "Three.js",
            "Bundled WebGL rendering support for 3D room viewers. Copyright (c) 2010-2026 three.js authors. "
            "Licensed under the MIT License.",
            "MIT",
        ),
    ]
    add_table(
        doc,
        ["Name in app", "What the app says", "License claimed"],
        [[a, b, c] for a, b, c in items],
    )

    doc.add_heading("How to validate (web only)", level=1)
    doc.add_paragraph(
        "For each row below: (1) open the official link, (2) confirm the license type matches what the app claims, "
        "(3) confirm commercial use is allowed (or note if it says non-commercial / research only), "
        "(4) write PASS, FAIL, or ASK LAWYER."
    )

    validation_rows = [
        [
            "Depth Anything V2",
            "Apache 2.0",
            "https://github.com/DepthAnything/Depth-Anything-V2",
            "Open LICENSE file and README. App uses the “Small” model only — README says Small is Apache; "
            "larger models are non-commercial. Confirm Small = OK for commercial.",
        ],
        [
            "GeoCalib",
            "Apache 2.0",
            "https://github.com/cvg/GeoCalib",
            "Open LICENSE — should be Apache 2.0, ETH Zurich copyright. Training-data fine print is a lawyer question.",
        ],
        [
            "MetalSplatter",
            "MIT",
            "https://github.com/scier/MetalSplatter",
            "Open LICENSE — MIT, Sean Cier copyright.",
        ],
        [
            "Firebase",
            "Apache 2.0",
            "https://github.com/firebase/firebase-ios-sdk",
            "Open LICENSE — Apache 2.0. Also check Google/Firebase terms of service for app use.",
        ],
        [
            "RTMDet / MMDetection",
            "Apache 2.0",
            "https://github.com/open-mmlab/mmdetection",
            "Open LICENSE — Apache 2.0, OpenMMLab. Model trained on COCO — optional: "
            "https://cocodataset.org/#termsofuse (annotations CC-BY 4.0). App does not mention COCO today.",
        ],
        [
            "Three.js",
            "MIT",
            "https://github.com/mrdoob/three.js",
            "Open LICENSE — MIT.",
        ],
        [
            "Apache License 2.0 (link in app)",
            "—",
            "https://www.apache.org/licenses/LICENSE-2.0",
            "Sanity-check that “View full license” for Apache items opens a real Apache 2.0 page.",
        ],
        [
            "MIT License (link in app)",
            "—",
            "https://opensource.org/licenses/MIT",
            "Sanity-check that “View full license” for MIT items opens a real MIT page.",
        ],
    ]
    add_table(
        doc,
        ["Component", "App claims", "Search here", "What to verify"],
        validation_rows,
    )

    doc.add_heading("Extra items (not shown in app — optional checks)", level=1)
    doc.add_paragraph(
        "These are used by the app but not listed on the Licenses screen. "
        "Only flag if you want a thorough review — not required for a quick pass."
    )
    extra = [
        [
            "spz-swift (splat file support)",
            "MIT",
            "https://github.com/scier/spz-swift",
            "Missing from in-app list — consider adding MIT credit.",
        ],
        [
            "Depth Anything training data (Hypersim)",
            "CC-BY-SA 3.0",
            "https://github.com/apple/ml-hypersim",
            "Not mentioned in app — lawyer question for commercial ship.",
        ],
    ]
    add_table(doc, ["Item", "License", "Search here", "Note"], extra)

    doc.add_heading("Simple checklist — fill in PASS / FAIL / ASK LAWYER", level=1)
    checks = [
        "I found the Licenses screen in the app.",
        "All six software names appear (Depth Anything, GeoCalib, MetalSplatter, Firebase, RTMDet, Three.js).",
        "“View full license” links work (open in Safari).",
        "Depth Anything — official site says Small model is Apache 2.0 (not non-commercial).",
        "GeoCalib — official LICENSE is Apache 2.0.",
        "MetalSplatter — official LICENSE is MIT.",
        "Firebase — official LICENSE is Apache 2.0.",
        "RTMDet / MMDetection — official LICENSE is Apache 2.0.",
        "Three.js — official LICENSE is MIT.",
        "Top notice (“non-commercial use only”) is acceptable for how the app will be sold.",
        "Anything unclear → escalated to a lawyer.",
    ]
    add_table(
        doc,
        ["#", "Check", "Result", "Notes"],
        [[str(i + 1), text, "", ""] for i, text in enumerate(checks)],
    )

    doc.add_heading("Reviewer sign-off", level=1)
    doc.add_paragraph("Name: ___________________________    Date: _______________")
    doc.add_paragraph("Overall:  PASS   /   FAIL   /   NEEDS LAWYER")
    doc.add_paragraph("Comments:")
    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_paragraph("")

    doc.add_paragraph(
        "Technical backup (optional): docs/MODEL_LICENSE_AUDIT.md in the project repo."
    )

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)
    print(f"Wrote {OUT_PATH}")


if __name__ == "__main__":
    main()
