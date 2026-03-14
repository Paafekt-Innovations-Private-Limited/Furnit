#!/usr/bin/env python3
"""
Generate ABAP AI Toolkit business-plan Word document.
Business language, non-technical. Patent opportunities in separate sections.
Output: abap-ai-toolkit/ABAP_AI_Toolkit_Business_Plan.docx
"""
from pathlib import Path

from docx import Document
from docx.shared import Pt
try:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    WD_ALIGN_PARAGRAPH = None  # fallback: no center align

SCRIPT_DIR = Path(__file__).resolve().parent
OUTPUT_DOCX = SCRIPT_DIR / "ABAP_AI_Toolkit_Business_Plan.docx"


def add_heading_styles(doc):
    """Ensure we have consistent heading styles."""
    # Use built-in styles; adjust if needed
    pass


def build_document():
    doc = Document()
    doc.add_paragraph()

    # ----- Title -----
    title = doc.add_paragraph()
    if WD_ALIGN_PARAGRAPH:
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ABAP AI Toolkit")
    run.bold = True
    run.font.size = Pt(22)
    run.font.name = "Calibri"

    subtitle = doc.add_paragraph()
    if WD_ALIGN_PARAGRAPH:
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("Business Overview & Patent Opportunities")
    sub_run.font.size = Pt(14)
    sub_run.font.name = "Calibri"
    doc.add_paragraph()

    # ----- Executive Summary -----
    doc.add_heading("Executive Summary", level=0)
    doc.add_paragraph(
        "The ABAP AI Toolkit is a software solution that helps enterprises move their "
        "custom SAP ECC systems to SAP S/4HANA Clean Core with less cost, fewer errors, "
        "and faster timelines. Instead of relying only on manual rewrites or generic code tools, "
        "the toolkit uses purpose-based matching, verified SAP compatibility data, and "
        "AI-assisted generation to recommend and generate modern replacements for legacy code. "
        "It is designed for CIOs, SAP program leads, and system integrators who need to reduce "
        "migration risk and accelerate their move to the cloud."
    )
    doc.add_paragraph()

    # ----- What We Offer -----
    doc.add_heading("What We Offer", level=0)
    doc.add_paragraph(
        "We offer an automated pipeline that takes a company's existing custom SAP (ABAP) code base "
        "and produces a clear migration path to S/4HANA Clean Core. The toolkit:"
    )
    bullets = [
        "Identifies which legacy interfaces and programs have official modern replacements.",
        "Maps old code to new APIs by business purpose, not just by name, so recommendations stay relevant even when SAP renames or restructures APIs.",
        "Modernizes syntax automatically where safe, and flags the rest for expert review.",
        "Builds a reusable knowledge base so that answers and examples stay consistent across teams and projects.",
        "Validates suggested and generated code against SAP's own documentation and rules, without requiring a full SAP system for every check.",
    ]
    for b in bullets:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(b)
    doc.add_paragraph()

    # ----- The Problem We Solve -----
    doc.add_heading("The Problem We Solve", level=0)
    doc.add_paragraph(
        "Many enterprises run critical business processes on older SAP ECC systems with large amounts "
        "of custom code. Moving to S/4HANA and Clean Core is mandatory for long-term support and "
        "innovation, but the migration is complex: thousands of programs, unclear replacement options, "
        "and a shortage of specialists. Manual analysis is slow and error-prone; generic translation "
        "tools do not understand SAP's release policies and API evolution. The result is delayed projects, "
        "budget overruns, and technical debt that persists after go-live."
    )
    doc.add_paragraph(
        "Our solution reduces that uncertainty by connecting legacy code to the right modern APIs, "
        "by purpose and with reference to SAP's official compatibility data, and by providing "
        "repeatable, documentable steps that support governance and audit."
    )
    doc.add_paragraph()

    # ----- How It Works (High-Level) -----
    doc.add_heading("How It Works", level=0)
    doc.add_paragraph(
        "The toolkit runs as a pipeline. A customer supplies an export of their custom SAP package "
        "(for example, from standard SAP export tools). The system then:"
    )
    steps = [
        ("Extract & organize", "The custom code is unpacked and organized by type (programs, classes, function modules, etc.)."),
        ("Scan & classify", "Legacy patterns and dependencies are detected and tagged for modernization or replacement."),
        ("Transform syntax", "Safe, rule-based updates are applied to bring syntax in line with current standards."),
        ("Find modern equivalents", "Legacy interfaces are matched to modern SAP APIs by what they do, not just by name, and checked against SAP's compatibility registry."),
        ("Build a knowledge base", "Code, rules, and mappings are indexed so that queries and AI-assisted generation can pull from verified, project-specific examples."),
        ("Generate & validate", "New or refactored code can be suggested and then checked against SAP documentation and cloud rules before handover to developers."),
    ]
    for label, text in steps:
        p = doc.add_paragraph()
        p.add_run(f"{label}. ").bold = True
        p.add_run(text)
    doc.add_paragraph(
        "The outcome is a migration roadmap, updated code where automation is applied, and a persistent "
        "knowledge base that improves consistency and speeds up future waves of migration."
    )
    doc.add_paragraph()

    # ----- Market Opportunity -----
    doc.add_heading("Market Opportunity", level=0)
    doc.add_paragraph(
        "SAP has set clear deadlines for S/4HANA adoption and Clean Core. Thousands of enterprises "
        "worldwide must migrate custom code. System integrators and internal IT teams are under "
        "pressure to deliver more with limited expert capacity. Demand for tools that reduce manual "
        "effort, improve accuracy, and provide traceable recommendations is strong. Our toolkit targets "
        "that gap: it is purpose-built for SAP migration, uses official SAP compatibility data where "
        "available, and can be offered as a licensed product, a managed service, or an embedded "
        "component in larger transformation programs."
    )
    doc.add_paragraph()

    # ========== PATENT OPPORTUNITIES (Separate Sections) ==========
    doc.add_heading("Patent Opportunities", level=0)
    doc.add_paragraph(
        "The following sections describe distinct innovations that may support patent protection. "
        "Each addresses a specific technical and business challenge in legacy-to-modern code migration "
        "and AI-assisted development. Legal advice should be sought to assess prior art and filing strategy."
    )
    doc.add_paragraph()

    # ----- Patent Opportunity 1: Semantic Twin Discovery -----
    doc.add_heading("1. Semantic Twin Discovery (Purpose-Based API Mapping)", level=1)
    doc.add_paragraph(
        "Today, many tools match old and new APIs by name or signature. In SAP's world, names and "
        "structures change frequently; one legacy interface may be replaced by several new ones, or "
        "several legacy components may collapse into one modern API. Name-based matching often fails "
        "or gives misleading results."
    )
    doc.add_paragraph(
        "Our approach is to match legacy code to modern APIs by business purpose: what the code "
        "\"does\" (e.g., \"post a journal entry,\" \"create a material document\") rather than what "
        "it is called. The system uses a purpose-oriented catalog of legacy-to-modern mappings and "
        "can represent one-to-one, one-to-many, many-to-one, and many-to-many relationships. When "
        "no direct replacement exists, it can recommend or generate wrapper designs that isolate "
        "legacy calls until SAP releases a full replacement."
    )
    doc.add_paragraph(
        "Patent opportunity: Methods and systems for determining replacement APIs for legacy code "
        "by semantic or purpose-based matching, including cardinality-aware mapping and automatic "
        "wrapper design, in the context of enterprise resource planning or similar domain-specific "
        "platforms."
    )
    doc.add_paragraph()

    # ----- Patent Opportunity 2: Cross-Lingual Skill Anchoring -----
    doc.add_heading("2. Cross-Lingual Skill Anchoring for Low-Resource Languages", level=1)
    doc.add_paragraph(
        "AI models are very good at popular programming languages (e.g., Java, Python) because "
        "they have been trained on huge amounts of public code. For niche or proprietary languages "
        "like SAP ABAP, public training data is scarce. As a result, models often \"hallucinate\" "
        "syntax or suggest constructs that do not exist or are obsolete."
    )
    doc.add_paragraph(
        "Our innovation is to use a well-known language (e.g., Java) as a \"skill anchor.\" We do "
        "not teach the model entirely new concepts; we provide a structured mapping that says: "
        "\"The skill you already know in Java corresponds to this construct in the target language.\" "
        "For example, the idea of \"get the size of a list\" is universal; we map the Java form "
        "to the correct, modern form in the target language. This mapping is used to guide "
        "code generation and transformation so that output stays aligned with real syntax and "
        "standards."
    )
    doc.add_paragraph(
        "Patent opportunity: Methods and systems for improving code generation or translation "
        "in a low-resource programming language by anchoring to a high-resource language through "
        "explicit, structured skill mappings (e.g., triple mappings: source language | legacy "
        "target | modern target), including use in AI-assisted migration and refactoring."
    )
    doc.add_paragraph()

    # ----- Patent Opportunity 3: Doc-Driven Validation Without Compiler -----
    doc.add_heading("3. Doc-Driven Validation Without a Compiler or Runtime", level=1)
    doc.add_paragraph(
        "In many environments, running the official compiler or runtime (e.g., the SAP kernel) is "
        "not possible: cloud build pipelines, developer laptops without SAP access, or third-party "
        "review tools. Without a compiler, it is hard to know if generated or migrated code is "
        "syntactically and stylistically correct."
    )
    doc.add_paragraph(
        "Our approach is to validate code using rules derived from the platform vendor's own "
        "documentation, style guides, and compatibility rules. We check for obsolete constructs, "
        "restricted language use, naming conventions, and structural patterns. The result is a "
        "quality score and a list of findings that approximate what a compiler or static "
        "checker would report, without requiring the actual compiler or runtime to be present."
    )
    doc.add_paragraph(
        "Patent opportunity: Methods and systems for validating source code in a proprietary or "
        "domain-specific language by applying documentation-derived and rule-based checks in "
        "environments where the official compiler or runtime is unavailable, including use in "
        "CI/CD, code review, and migration tooling."
    )
    doc.add_paragraph()

    # ----- Patent Opportunity 4: Integrated Pipeline (Optional) -----
    doc.add_heading("4. Integrated Migration Pipeline (Combined Innovation)", level=1)
    doc.add_paragraph(
        "A further patent opportunity lies in the combination of the above: an integrated pipeline "
        "that (a) extracts and classifies legacy code, (b) applies purpose-based twin discovery "
        "with compatibility verification against the vendor's registry, (c) indexes and retrieves "
        "knowledge for AI-assisted generation, and (d) validates outputs using documentation-based "
        "rules. The orchestration, data flow, and use of a single knowledge base for both human "
        "queries and machine-generated code can be claimed as a system and method for automated "
        "legacy-to-modern migration in enterprise software."
    )
    doc.add_paragraph()

    # ----- Competitive Advantage -----
    doc.add_heading("Competitive Advantage", level=0)
    doc.add_paragraph(
        "Our differentiators include: (1) purpose-based rather than name-based API matching; "
        "(2) use of SAP's public compatibility data to ground recommendations; (3) skill anchoring "
        "to improve AI output quality for a low-resource language; (4) validation without "
        "dependency on a full SAP system; and (5) a reusable, shareable knowledge base that "
        "improves over time. These innovations support both product differentiation and "
        "potential patent protection as outlined above."
    )
    doc.add_paragraph()

    # ----- Next Steps -----
    doc.add_heading("Next Steps", level=0)
    doc.add_paragraph(
        "Recommended next steps for the business plan and IP strategy: (1) Validate market size "
        "and willingness to pay with target customers and partners. (2) Commission a prior-art "
        "search and patentability opinion for the innovations in Sections 1–4. (3) Define "
        "productization and licensing options (e.g., SaaS, on-premise, white-label). (4) Identify "
        "pilot customers for case studies and referenceability."
    )
    doc.add_paragraph()

    # ----- Footer note -----
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Document generated for business planning purposes. ")
    p.add_run("Technical details and implementation are available in the ABAP AI Toolkit repository and documentation.")
    p.paragraph_format.space_before = Pt(12)

    return doc


def main():
    doc = build_document()
    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_DOCX))
    print(f"Saved: {OUTPUT_DOCX}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
