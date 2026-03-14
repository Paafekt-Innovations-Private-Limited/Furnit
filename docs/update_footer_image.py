#!/usr/bin/env python3
"""
Replace footer image in an existing letterhead .docx with docs/letterhead/bottom.png.
Use when bottom.png has been fixed and you want to update the docx without re-running
add_letterhead from scratch.
"""
from pathlib import Path

DOCS_DIR = Path(__file__).resolve().parent
LETTERHEAD_BOTTOM = DOCS_DIR / "letterhead" / "bottom.png"


def update_footer_image(docx_path: Path) -> None:
    from docx import Document
    from docx.shared import Inches

    if not docx_path.exists():
        raise FileNotFoundError(docx_path)
    if not LETTERHEAD_BOTTOM.exists():
        raise FileNotFoundError(LETTERHEAD_BOTTOM)

    doc = Document(str(docx_path))
    footer_width = Inches(6.5)

    for section in doc.sections:
        footer = section.footer
        if not footer.paragraphs:
            continue
        first_para = footer.paragraphs[0]
        first_para.clear()
        new_run = first_para.add_run()
        new_run.add_picture(str(LETTERHEAD_BOTTOM), width=footer_width)

    doc.save(str(docx_path))
    print(f"Updated footer image in {docx_path} using {LETTERHEAD_BOTTOM}")


if __name__ == "__main__":
    import sys
    path = Path(sys.argv[1]) if len(sys.argv) > 1 else DOCS_DIR / "NDA - Paafekt v02_0903 (1) - DIN11594139 - letterhead.docx"
    update_footer_image(path)
